import os
import json
import io
import asyncio

import numpy as np
from typing import List, Dict, Any, Optional
from fastapi import FastAPI, HTTPException, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from openai import AsyncOpenAI
from dotenv import load_dotenv

import pdfplumber
import docx
from odf import text, teletype
from odf.opendocument import load

load_dotenv()

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

client = AsyncOpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# Models
class ParseResponse(BaseModel):
    text: str

class AnalyzeRequest(BaseModel):
    resume_text: str
    job_description: str

class ResumeData(BaseModel):
    candidate_name: Optional[str] = None
    skills: List[str]
    experience: List[str]
    education: List[str]

class JobData(BaseModel):
    company_name: Optional[str] = None
    job_role: Optional[str] = None
    required_skills: List[str]
    preferred_skills: List[str]
    responsibilities: List[str]

class AnalyzeResponse(BaseModel):
    match_score: int
    skill_overlap: List[str]
    missing_skills: List[str]
    improved_bullets: List[str]
    ats_keywords: Dict[str, Any]
    summary: List[str]
    candidate_name: Optional[str] = None
    company_name: Optional[str] = None
    job_role: Optional[str] = None

class PitchRequest(BaseModel):
    resume_text: str
    job_description: str
    match_score: int
    skill_overlap: List[str]
    length: str # "short" | "extended"
    tone: str # "formal" | "casual"
    job_role: Optional[str] = None
    company_name: Optional[str] = None

class PitchResponse(BaseModel):
    pitch: str

# Utils
def load_prompt(filename: str, **kwargs) -> str:
    path = os.path.join("backend/prompts", filename)
    if not os.path.exists(path):
        # Fallback for if we are running from inside backend dir
        path = os.path.join("prompts", filename)
    with open(path, "r") as f:
        prompt = f.read()
    for key, value in kwargs.items():
        prompt = prompt.replace(f"{{{{{key}}}}}", str(value))
    return prompt

async def get_completion(prompt: str, json_mode: bool = False) -> str:
    response = await client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.2,
        response_format={"type": "json_object"} if json_mode else None
    )
    return response.choices[0].message.content

async def get_embeddings(texts: List[str]) -> List[List[float]]:
    if not texts:
        return []
    # OpenAI supports up to 2048 inputs per request for embeddings
    response = await client.embeddings.create(
        input=texts,
        model="text-embedding-3-small"
    )
    return [d.embedding for d in response.data]

async def get_embedding(text: str) -> List[float]:
    res = await get_embeddings([text])
    return res[0]

def cosine_similarity(a, b):
    return np.dot(a, b) / (np.linalg.norm(a) * np.linalg.norm(b))

async def calculate_score(resume_skills: List[str], job_data: JobData) -> int:
    if not resume_skills:
        return 0
    if not job_data.required_skills and not job_data.preferred_skills:
        return 0

    # Step 1: Generate all embeddings in batches
    resume_embs = await get_embeddings(resume_skills)
    
    # Helper for BestMatchScore
    def get_best_match(job_skill_emb, res_embs):
        if not res_embs:
            return 0.0
        similarities = [cosine_similarity(job_skill_emb, res_emb) for res_emb in res_embs]
        return max(similarities)

    # Step 2-4: Required Skills
    req_score = 0.0
    if job_data.required_skills:
        req_embs = await get_embeddings(job_data.required_skills)
        req_best_matches = [get_best_match(emb, resume_embs) for emb in req_embs]
        
        # Step 3: Coverage (Threshold >= 0.55)
        matched_required = [m for m in req_best_matches if m >= 0.55]
        req_coverage = len(matched_required) / len(job_data.required_skills)
        
        # Step 4: Strength (Mean of matched skills)
        req_strength = np.mean(matched_required) if matched_required else 0.0
        
        # Step 5 (part): Required Score
        req_score = (req_coverage * 0.6) + (req_strength * 0.4)
    else:
        # If no required skills, we give full marks for this section
        req_score = 1.0

    # Step 5: Preferred Skills
    pref_strength = 0.0
    if job_data.preferred_skills:
        pref_embs = await get_embeddings(job_data.preferred_skills)
        pref_best_matches = [get_best_match(emb, resume_embs) for emb in pref_embs]
        pref_strength = np.mean(pref_best_matches) if pref_best_matches else 0.0
    else:
        # If no preferred skills, we match it to req_score so it doesn't affect the weighted average
        pref_strength = req_score

    # Step 6: Final Weighted Score
    # Weighting: Required 70%, Preferred 30%
    if not job_data.preferred_skills:
        final_score = req_score
    elif not job_data.required_skills:
        final_score = pref_strength
    else:
        final_score = (req_score * 0.7) + (pref_strength * 0.3)
    
    return int(round(min(max(final_score, 0), 1) * 100))

def normalize_text(text: str) -> str:
    # Normalize whitespace
    text = "\n".join([line.strip() for line in text.split("\n")])
    # Remove repeated empty lines
    import re
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()

def ensure_list_of_strings(data: Any) -> List[str]:
    if not isinstance(data, list):
        return []
    result = []
    for item in data:
        if isinstance(item, str):
            result.append(item)
        elif isinstance(item, dict):
            # Flatten dictionary into a single string
            result.append(" ".join(str(v) for v in item.values() if v))
        else:
            result.append(str(item))
    return result

@app.post("/parse-resume", response_model=ParseResponse)
async def parse_resume(file: UploadFile = File(...)):
    filename = file.filename.lower()
    content_type = file.content_type
    
    try:
        content = await file.read()
        if not content:
            raise HTTPException(status_code=400, detail="Empty file")

        extracted_text = ""
        
        if filename.endswith(".pdf") or content_type == "application/pdf":
            try:
                with pdfplumber.open(io.BytesIO(content)) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            extracted_text += page_text + "\n"
            except Exception as e:
                print(f"pdfplumber error: {e}")
                raise HTTPException(status_code=400, detail="Unreadable PDF document")
            
            if len(extracted_text.strip()) < 50: # Lowered threshold slightly
                 raise HTTPException(status_code=400, detail="This PDF appears to be a scanned image or has very little text. Please upload a text-based PDF or paste text manually.")

        elif filename.endswith(".docx") or content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            try:
                doc = docx.Document(io.BytesIO(content))
                extracted_text = "\n".join([para.text for para in doc.paragraphs])
            except Exception as e:
                print(f"python-docx error: {e}")
                raise HTTPException(status_code=400, detail="Unreadable DOCX document")
            
        elif filename.endswith(".odt") or content_type == "application/vnd.oasis.opendocument.text":
            try:
                odt_doc = load(io.BytesIO(content))
                paragraphs = odt_doc.getElementsByType(text.P)
                extracted_text = "\n".join([teletype.extractText(p) for p in paragraphs])
            except Exception as e:
                print(f"odfpy error: {e}")
                raise HTTPException(status_code=400, detail="Unreadable ODT document")
            
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type. Please upload a PDF, DOCX, or ODT file.")

        if not extracted_text.strip():
            raise HTTPException(status_code=400, detail="Could not extract text from the document.")

        return ParseResponse(text=normalize_text(extracted_text))

    except HTTPException:
        raise
    except Exception as e:
        print(f"Parsing error: {e}")
        raise HTTPException(status_code=500, detail="Internal server error during parsing")

@app.post("/parse-jd", response_model=ParseResponse)
async def parse_jd(file: UploadFile = File(...)):
    filename = file.filename.lower()
    content_type = file.content_type
    
    try:
        content = await file.read()
        if not content:
            raise HTTPException(status_code=400, detail="Empty file")

        extracted_text = ""
        
        if filename.endswith(".pdf") or content_type == "application/pdf":
            try:
                with pdfplumber.open(io.BytesIO(content)) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            extracted_text += page_text + "\n"
            except Exception as e:
                print(f"pdfplumber error: {e}")
                raise HTTPException(status_code=400, detail="Unreadable JD PDF document")
            
            if len(extracted_text.strip()) < 50:
                 raise HTTPException(status_code=400, detail="This JD PDF appears to be a scanned image or has very little text. Please upload a text-based PDF or paste text manually.")

        elif filename.endswith(".docx") or content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            try:
                doc = docx.Document(io.BytesIO(content))
                extracted_text = "\n".join([para.text for para in doc.paragraphs])
            except Exception as e:
                print(f"python-docx error: {e}")
                raise HTTPException(status_code=400, detail="Unreadable JD DOCX document")
            
        elif filename.endswith(".odt") or content_type == "application/vnd.oasis.opendocument.text":
            try:
                odt_doc = load(io.BytesIO(content))
                paragraphs = odt_doc.getElementsByType(text.P)
                extracted_text = "\n".join([teletype.extractText(p) for p in paragraphs])
            except Exception as e:
                print(f"odfpy error: {e}")
                raise HTTPException(status_code=400, detail="Unreadable JD ODT document")
            
        else:
            raise HTTPException(status_code=400, detail="Unsupported JD file type. Please upload a PDF, DOCX, or ODT file.")

        if not extracted_text.strip():
            raise HTTPException(status_code=400, detail="Could not extract text from the JD document.")

        return ParseResponse(text=normalize_text(extracted_text))

    except HTTPException:
        raise
    except Exception as e:
        print(f"JD Parsing error: {e}")
        raise HTTPException(status_code=500, detail="Internal server error during JD parsing")

@app.post("/analyze", response_model=AnalyzeResponse)
async def analyze(request: AnalyzeRequest):
    if not request.resume_text.strip() or not request.job_description.strip():
        raise HTTPException(status_code=400, detail="Resume and Job Description are required")
    try:
        # 1. Parse Data (Phase 1 Parallel)
        resume_prompt = load_prompt("parse_resume.txt", resume_text=request.resume_text)
        job_prompt = load_prompt("parse_job.txt", job_description=request.job_description)
        
        resume_raw, job_raw = await asyncio.gather(
            get_completion(resume_prompt, json_mode=True),
            get_completion(job_prompt, json_mode=True)
        )
        
        resume_json = json.loads(resume_raw)
        
        # Ensure experience and education are lists of strings
        resume_json['skills'] = ensure_list_of_strings(resume_json.get('skills', []))
        resume_json['experience'] = ensure_list_of_strings(resume_json.get('experience', []))
        resume_json['education'] = ensure_list_of_strings(resume_json.get('education', []))
        
        resume_data = ResumeData(**resume_json)

        job_json = json.loads(job_raw)
        
        # Also ensure job data lists are strings
        job_json['required_skills'] = ensure_list_of_strings(job_json.get('required_skills', []))
        job_json['preferred_skills'] = ensure_list_of_strings(job_json.get('preferred_skills', []))
        job_json['responsibilities'] = ensure_list_of_strings(job_json.get('responsibilities', []))
        
        job_data = JobData(**job_json)

        # 2. Match Score (Deterministic via Embeddings)
        match_score = await calculate_score(resume_data.skills, job_data)

        # 3-6. Analysis (Phase 2 Parallel)
        gap_prompt = load_prompt("skill_gap.txt", skills=resume_data.skills, job_requirements=job_data.required_skills)
        rewrite_prompt = load_prompt("resume_rewrite.txt", bullets=resume_data.experience, job_description=request.job_description)
        ats_prompt = load_prompt("ats_keywords.txt", job_description=request.job_description, resume_text=request.resume_text)
        summary_prompt = load_prompt("final_summary.txt", resume_text=request.resume_text, job_description=request.job_description)
        
        gap_raw, rewrite_raw, ats_raw, summary_raw = await asyncio.gather(
            get_completion(gap_prompt, json_mode=True),
            get_completion(rewrite_prompt),
            get_completion(ats_prompt, json_mode=True),
            get_completion(summary_prompt)
        )

        
        gap_data = json.loads(gap_raw)

        # Improved Bullets
        try:
            improved_bullets = json.loads(rewrite_raw)
            if not isinstance(improved_bullets, list):
                improved_bullets = [rewrite_raw]
        except:
            improved_bullets = [b.strip("- ") for b in rewrite_raw.split("\n") if b.strip()]

        # ATS Keywords
        ats_keywords = json.loads(ats_raw)

        # Final Summary
        summary = [s.strip("- ") for s in summary_raw.split("\n") if s.strip()]

        return AnalyzeResponse(
            match_score=match_score,
            skill_overlap=gap_data.get("strong_matches", []) + gap_data.get("partial_matches", []),
            missing_skills=gap_data.get("missing_skills", []),
            improved_bullets=improved_bullets,
            ats_keywords=ats_keywords,
            summary=summary[:4], # Strictly 4 points
            candidate_name=resume_data.candidate_name,
            company_name=job_data.company_name,
            job_role=job_data.job_role
        )

    except Exception as e:
        print(f"Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/generate-pitch", response_model=PitchResponse)
async def generate_pitch(request: PitchRequest):
    try:
        # Enforce Tone Rules: If length is "extended", tone is formal
        effective_tone = request.tone
        if request.length == "extended":
            effective_tone = "formal"
        
        # Mapping for Prompt
        length_mapping = {
            "short": "3-5 concise sentences",
            "extended": "Structured, up to one page"
        }
        tone_mapping = {
            "formal": "Professional, direct",
            "casual": "Conversational, but respectful"
        }

        prompt = load_prompt(
            "personal_pitch.txt",
            resume_highlights=request.resume_text[:2000], # Limit to avoid token issues
            job_description=request.job_description[:2000],
            job_role=request.job_role or "Target Role",
            company_name=request.company_name or "Target Company",
            skill_overlap=", ".join(request.skill_overlap),
            match_score=request.match_score,
            length=length_mapping.get(request.length, request.length),
            tone=tone_mapping.get(effective_tone, effective_tone)
        )

        pitch = await get_completion(prompt)
        return PitchResponse(pitch=pitch.strip())

    except Exception as e:
        print(f"Pitch generation error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
